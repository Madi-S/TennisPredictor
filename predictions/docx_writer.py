from docx import Document


bio = ['Name', 'Age', 'Country', 'Ranking', 'RankingPeak',
       'Points', 'PrizeMoney', 'TotalMatches', 'Winrate%']
order = ['Outcome', 'Odds',  'Explanation']


class DOCXWriter:

    def __init__(self, filename: str):
        '''
        Initialize DOCX Writer by creating DOCX with given filename

        :param filename: `str` for the DOCX file name
        :return: returns nothing
        '''
        self._filename = filename + '.docx'
        self._d = Document()

    def write(self, betting_tips: dict, stats: dict, data: dict):
        '''
        Save given data to .docx file

        :param betting_tips: `dict` of betting tips, odds, bets tendency, players' names
        :param stats: `dict` containing detailed statistics of each player
        :param data: `dict` containing time, h2h, conclusion and points scored for each player
        :return: returns nothing
        '''
        players = betting_tips['Players']
        p1, p2 = players

        self._d.add_heading(f'{p1} ({betting_tips["Odds"][p1]}) vs {p2} ({betting_tips["Odds"][p2]})', 0)
        self._d.add_heading(f'{betting_tips["Tournament"].replace(".",". ")}', 0)
        self._d.add_heading(f'{data["Time"]}\n', 0)

        for p in players:
            for o in bio:
                self._d.add_paragraph(f'{o}: {stats[p][o]}')
            self._d.add_paragraph(f'Bets tendency on winner - {p}: {betting_tips["BetsTendency"][p]}\n')
        
        self._d.add_paragraph(f'Bets tendency on total over: {betting_tips["BetsTendency"]["TotalOver"]}')
        self._d.add_paragraph(f'Bets tendency on total under: {betting_tips["BetsTendency"]["TotalUnder"]}')

        self._d.add_paragraph(f'Head to heads: {data["H2H"]}')

        p = self._d.add_paragraph('\n')
        p.add_run('Top Betting tips:').bold = True

        for b in betting_tips['Predictions']:
            for o in order:
                self._d.add_paragraph(f'{o}: {b[o]}')
            self._d.add_paragraph('\n')

        p = self._d.add_paragraph('\n')
        p.add_run(f'Conclusion: {p1} scored {data["Points"][p1]} and {p2} scored {data["Points"][p2]}\n{data["Conclusion"]}').bold = True

        self._d.save(self._filename)


def main():
    d = DOCXWriter('TennisPredictions')
    betting_tips = {'Players': ['Djokovic N', 'Nadal R'], 'Tournament': 'Теннис. ATP. Анталья. Квалификация',
                    'Odds': {'Djokovic N': 1.68, 'Nadal R': 2.3}, 'BetsTendency':
                    {'Djokovic N': 234.7, 'Nadal R': 231.9,
                     'TotalOver': 81.2, 'TotalUnder': 43.5},
                    'PastResults': {'Djokovic N': ['-', '+', '-', '+', '-', '+'], 'Nadal R': ['-', '+', '+', '+', '+', '+']},
                    'Predictions': [{'Outcome': 'Handicap2 by games (4.5)', 'Odds': 1.64, 
                    'Explanation': 'Behincher will play against Orlov.There were no faces.Behincher is more experienced and cunning player onclay today in Antalya will play on hard in cool weather. Here the German can smash Orel with -6.5 handicap.The bet is risky.Orlov I liked the game several times I saw him in action and the ball flies into the cornerand feeds are even and bold access to the net.In general, kkk for me, then Bechinger should win reliably,but a few views of Orlov give doubts one hundred Bnchinger', 
                    'ExpertProfit % ': 10.74},
                     {'Outcome': 'WINNER 2', 'Odds': 2.29, 
                     'Explanation': "Friends, my next prediction for Orlov's victory!Let the Bachinger be much moreexperienced and skilled, but I believe that his time has come to please the youngsters, and the Bachinger hasa good serve, not very good at the back line and very bad at the net.Even if Orlov is of the level of ITF, buthe is skating on all his rivals, closes 2-0, unlike his counterpart, he is played, he took a short pause beforeAntalya, I think that this tournament was focused on.", 
                     'ExpertProfit%': 14.72}]}
    
    stats = {'Djokovic N': {'Name': 'Novak Djokovic', 'Age': 33, 'Ranking': 1, 'RankingPeak': 1, 'Points': 12000, 'PrizeMoney': 12400003, 'TotalMatches': 1108, 'Winrate%': 81},
    'Nadal R': {'Name': 'Rafael Nadal', 'Age': 34, 'Ranking': 2, 'RankingPeak': 1, 'Points': 11000, 'PrizeMoney': 10400003, 'TotalMatches': 1143, 'Winrate%': 84}}
    data = {'Time': '11:15', 'Points': {'Djokovic N': 2391.333336, 'Nadal R': 1998.432}, 'Conclusion': 'Considering various in-games statistics, past results and h2h, Djokovic should win', 'Probability': '74%', 'H2H': None}
    d.write(betting_tips, stats, data)


if __name__ == '__main__':
    main()
