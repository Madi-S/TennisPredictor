from docx import Document
from random import choice


easy_game = (
    '{w} excelled {l} in many stats. Moreover, {w} in a good shape, comparing to {l}. To sum up, {w} should win here confidently',
    'Considering players\' statistics, {w} is much better than {l}. What is more, many betting experts suggest betting on {w}. Hence, my prediction - winner {w}',
    'Today\'s match should be a cakewalk for {w}. Taking into consideration that he scored more points than {l}, {w} has to take victory over {l}',
    '{w} is an apparent favourite here. Tennis experts say that {l} has little chances here. {w} and {l} are on different levels, so {w} should win today'
)
hard_game = (
    'This will be a tight fight between {l} and {w}. Some tennis portals presume that {w} will triumph over {l}. Nonetheless, others consider {l} as favourite. Final prediction: small bet on {w}',
    '{w} and {l} are relatively equal tennis players. Of course, {w} surpass {l} in some indicators. However, {l} also shows some good tennis. Accordingly, a risky match here but {w} has a small advatage over {l}',
    '{l} vs {w} - a clash of two big tennis figures. Players have scored approximately equal amount of points but {w} scored a little bit more. Thus and so, small bet on {w} is a good choice',
    'An even match for today: {l} vs {w}. {l} and {w} are in good shape and anyone can win today as tennis experts assume. Nevertheless, {w} has scored more points than {l} and this is a sign for a small bet on {w}',
)


class DOCXWriter:

    def __init__(self, filename: str):
        self._filename = filename + '.docx'
        self._d = Document()

    def write_(self, w: str, l: str, prob: int, time: str, tournament: str, easy: bool):
        self._d.add_heading(f'{w} vs {l}', 0)
        self._d.add_heading(tournament, 0)
        self._d.add_heading(time, 0)

        if easy == True:
            self._d.add_paragraph(choice(easy_game).format(w=w, l=l))
        else:
            self._d.add_paragraph(choice(hard_game).format(w=w, l=l))
        self._d.add_paragraph(f'Probability: {prob}%')
        self._d.add_paragraph('\n\n')
        self._d.save(self._filename)

    def write(self, p1, p2, p1_oddds, p2_odds, w, prob, time, tournament):
        self._d.add_heading(f'{p1} vs {p2}', 0)
        self._d.add_heading(f'Tournament: {tournament}', 0)
        if not time:
            self._d.add_heading(f'Time is not specified yet')
        else:
            self._d.add_heading(f'Time: {time}', 0)
        self._d.add_paragraph(f'{p1} odds: {p1_oddds}')
        self._d.add_paragraph(f'{p2} odds: {p2_odds}')
        self._d.add_paragraph(f'Winner: {w}')
        self._d.add_paragraph(f'Probability: {prob}%')
        self._d.add_paragraph('\n\n\n')
        self._d.save(self._filename)


if __name__ == '__main__':
    d = DOCXWriter('test')
    for _ in range(5):
        d.write('TEST_1', 'TEST_2', 'TEST_ODDS_2', 'TEST_ODDS_1', 'TEST_WINNER', 'TEST_PROBABILITY', 'TEST_TIME', 'TEST_TOURNAMENT')
