from docx import Document
from random import choice



class DOCXWriter:

    def __init__(self, filename: str):
        self._filename = filename + '.docx'
        self._d = Document()
        with open(self._filename, 'w') as f:
            f.write('')

    def write_(self, w: str, l: str, prob: int, time: str, tournament: str, easy: bool):
        self._d.add_heading(f'{w} vs {l}', 0)
        self._d.add_heading(tournament, 0)
        self._d.add_heading(time, 0)

        if easy == True:
            self._d.add_paragraph(choice(easy_game).format(w=w, l=l))
        else:
            self._d.add_paragraph(choice(hard_game).format(w=w, l=l))
        self._d.add_paragraph(f'Probability: {prob}%')
        self._d.add_paragraph('\n\n')
        self._d.save(self._filename)

    def write(self, p1, p2, p1_oddds, p2_odds, w, prob, time, tournament):
        self._d.add_heading(f'{p1} vs {p2}', 0)
        self._d.add_heading(f'Tournament: {tournament}', 0)
        if not time:
            self._d.add_heading(f'Time is not specified yet')
        else:
            self._d.add_heading(f'Time: {time}', 0)
        self._d.add_paragraph(f'{p1} odds: {p1_oddds}')
        self._d.add_paragraph(f'{p2} odds: {p2_odds}')
        self._d.add_paragraph(f'Winner: {w}')
        self._d.add_paragraph(f'Probability: {prob}%')
        self._d.add_paragraph('\n\n\n')
        self._d.save(self._filename)
